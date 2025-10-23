"""
Enhanced AI-Powered Resume Parser with Comprehensive Error Handling

This module provides production-ready resume parsing with:
1. Robust error handling
2. Validation at every step
3. Detailed logging
4. Fallback mechanisms
5. Performance monitoring
"""

import os
import re
import json
import logging
import time
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from pathlib import Path
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import spacy
from openai import OpenAI
import pytesseract
from PIL import Image
import io
import phonenumbers
from email_validator import validate_email, EmailNotValidError
from .skill_ranker import rank_and_classify_skills

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ResumeParsingError(Exception):
    """Base exception for resume parsing errors"""
    pass


class FileValidationError(ResumeParsingError):
    """Raised when file validation fails"""
    pass


class TextExtractionError(ResumeParsingError):
    """Raised when text extraction fails"""
    pass


class LLMEnhancementError(ResumeParsingError):
    """Raised when LLM enhancement fails"""
    pass


class EnhancedResumeParser:
    """
    Production-ready resume parser with comprehensive error handling.
    """
    
    # Supported file types
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    
    # Max file size (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Min text length for valid resume
    MIN_TEXT_LENGTH = 100
    
    def __init__(self, openai_api_key: Optional[str] = None):
        """Initialize the parser with error handling."""
        self.start_time = None
        self.errors = []
        self.warnings = []
        
        # Load spaCy model
        try:
            try:
                self.nlp = spacy.load("en_core_web_lg")
                logger.info("Loaded spaCy model: en_core_web_lg")
            except OSError:
                logger.warning("Large spaCy model not found, using small model")
                self.nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            logger.error(f"Failed to load spaCy model: {e}")
            self.nlp = None
            self.warnings.append("NER functionality disabled")
        
        # Initialize OpenAI client
        try:
            api_key = openai_api_key or os.getenv('OPENAI_API_KEY')
            if api_key:
                self.openai_client = OpenAI(api_key=api_key)
                logger.info("OpenAI client initialized")
            else:
                self.openai_client = None
                self.warnings.append("OpenAI API key not provided - LLM enhancement disabled")
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI client: {e}")
            self.openai_client = None
            self.warnings.append("LLM enhancement disabled due to initialization error")
    
    def parse(self, file_path: str) -> Dict:
        """
        Main parsing function with comprehensive error handling.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary with structured resume data
            
        Raises:
            FileValidationError: If file validation fails
            TextExtractionError: If text extraction fails
            ResumeParsingError: For other parsing errors
        """
        self.start_time = time.time()
        self.errors = []
        self.warnings = []
        
        logger.info(f"Starting to parse resume: {file_path}")
        
        try:
            # Step 1: Validate file
            self._validate_file(file_path)
            
            # Step 2: Extract text
            text, is_scanned = self._extract_text_with_fallback(file_path)
            
            # Step 3: Validate extracted text
            self._validate_text(text)
            
            # Step 4: Extract entities with spaCy (if available)
            entities = self._extract_entities_safely(text)
            
            # Step 5: Enhance with LLM (if available)
            structured_data = self._enhance_with_llm_safely(text, entities)
            
            # Step 6: Post-process and validate
            final_data = self._post_process_and_validate(structured_data, text)
            
            # Step 7: Add metadata
            final_data['metadata'] = self._create_metadata(file_path, is_scanned)
            
            logger.info(f"Successfully parsed resume in {time.time() - self.start_time:.2f}s")
            
            return final_data
            
        except Exception as e:
            logger.error(f"Resume parsing failed: {e}")
            self.errors.append(str(e))
            
            # Return error response
            return {
                'success': False,
                'error': str(e),
                'errors': self.errors,
                'warnings': self.warnings,
                'metadata': {
                    'file_path': file_path,
                    'parsed_at': datetime.now().isoformat(),
                    'parsing_time_seconds': time.time() - self.start_time if self.start_time else 0,
                    'parser_version': '2.0.0'
                }
            }
    
    def _validate_file(self, file_path: str):
        """Validate file exists, size, and extension."""
        # Check file exists
        if not os.path.exists(file_path):
            raise FileValidationError(f"File not found: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise FileValidationError("File is empty")
        
        if file_size > self.MAX_FILE_SIZE:
            raise FileValidationError(
                f"File too large: {file_size / 1024 / 1024:.2f}MB "
                f"(max: {self.MAX_FILE_SIZE / 1024 / 1024:.0f}MB)"
            )
        
        # Check file extension
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise FileValidationError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        logger.info(f"File validation passed: {file_size / 1024:.2f}KB, {ext}")
    
    def _extract_text_with_fallback(self, file_path: str) -> Tuple[str, bool]:
        """Extract text with multiple fallback methods."""
        ext = Path(file_path).suffix.lower()
        is_scanned = False
        
        try:
            if ext == '.docx':
                text = self._extract_from_docx(file_path)
            elif ext == '.pdf':
                text = self._extract_from_pdf_with_fallback(file_path)
                
                # Check if PDF is scanned (very little text)
                if len(text.strip()) < 100:
                    logger.warning("PDF appears to be scanned, attempting OCR")
                    text = self._extract_with_ocr(file_path)
                    is_scanned = True
            else:
                raise TextExtractionError(f"Unsupported file type: {ext}")
            
            logger.info(f"Text extracted: {len(text)} characters")
            return text, is_scanned
            
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text: {e}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            text = "\n".join(paragraphs)
            
            if not text.strip():
                raise TextExtractionError("No text found in DOCX file")
            
            return text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _extract_from_pdf_with_fallback(self, file_path: str) -> str:
        """Extract text from PDF with fallback methods."""
        text = ""
        
        # Method 1: Try pdfplumber (best for layout-aware extraction)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if len(text.strip()) > 100:
                logger.info("PDF extraction successful with pdfplumber")
                return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        
        # Method 2: Try PyMuPDF
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            
            if len(text.strip()) > 100:
                logger.info("PDF extraction successful with PyMuPDF")
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF failed: {e}")
        
        # If both methods failed or extracted very little text
        if len(text.strip()) < 100:
            raise TextExtractionError("Failed to extract sufficient text from PDF")
        
        return text
    
    def _extract_with_ocr(self, file_path: str) -> str:
        """Extract text using OCR (Tesseract)."""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(min(len(doc), 10)):  # Limit to 10 pages
                page = doc[page_num]
                pix = page.get_pixmap()
                img = Image.open(io.BytesIO(pix.tobytes()))
                
                # Run OCR
                page_text = pytesseract.image_to_string(img)
                text += page_text + "\n"
            
            doc.close()
            
            if len(text.strip()) < 100:
                raise TextExtractionError("OCR extracted insufficient text")
            
            logger.info(f"OCR extraction successful: {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise TextExtractionError(f"OCR failed: {e}")
    
    def _validate_text(self, text: str):
        """Validate extracted text."""
        if not text or not text.strip():
            raise TextExtractionError("No text extracted from file")
        
        if len(text.strip()) < self.MIN_TEXT_LENGTH:
            raise TextExtractionError(
                f"Insufficient text extracted: {len(text)} characters "
                f"(minimum: {self.MIN_TEXT_LENGTH})"
            )
        
        logger.info(f"Text validation passed: {len(text)} characters")
    
    def _extract_entities_safely(self, text: str) -> Dict:
        """Extract entities with error handling."""
        if not self.nlp:
            logger.warning("spaCy not available, skipping NER")
            return self._extract_entities_with_regex(text)
        
        try:
            doc = self.nlp(text[:10000])  # Limit to first 10K chars for performance
            
            entities = {
                "persons": [],
                "organizations": [],
                "locations": [],
                "dates": [],
                "emails": [],
                "phones": [],
                "urls": []
            }
            
            # Extract named entities
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    entities["persons"].append(ent.text)
                elif ent.label_ == "ORG":
                    entities["organizations"].append(ent.text)
                elif ent.label_ in ["GPE", "LOC"]:
                    entities["locations"].append(ent.text)
                elif ent.label_ == "DATE":
                    entities["dates"].append(ent.text)
            
            # Add regex-based extraction
            regex_entities = self._extract_entities_with_regex(text)
            entities["emails"] = regex_entities["emails"]
            entities["phones"] = regex_entities["phones"]
            entities["urls"] = regex_entities["urls"]
            
            logger.info(f"Entities extracted: {sum(len(v) for v in entities.values())} total")
            return entities
            
        except Exception as e:
            logger.error(f"NER failed: {e}")
            self.warnings.append("Named Entity Recognition failed, using regex fallback")
            return self._extract_entities_with_regex(text)
    
    def _extract_entities_with_regex(self, text: str) -> Dict:
        """Extract entities using regex patterns (fallback)."""
        entities = {
            "persons": [],
            "organizations": [],
            "locations": [],
            "dates": [],
            "emails": [],
            "phones": [],
            "urls": []
        }
        
        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        entities["emails"] = list(set(re.findall(email_pattern, text)))
        
        # Extract phone numbers
        phone_patterns = [
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
        ]
        for pattern in phone_patterns:
            entities["phones"].extend(re.findall(pattern, text))
        entities["phones"] = list(set(entities["phones"]))
        
        # Extract URLs
        url_pattern = r'https?://(?:www\.)?[-a-zA-Z0-9@:%._\+~#=]{1,256}\.[a-zA-Z0-9()]{1,6}\b(?:[-a-zA-Z0-9()@:%_\+.~#?&/=]*)'
        entities["urls"] = list(set(re.findall(url_pattern, text)))
        
        return entities
    
    def _enhance_with_llm_safely(self, text: str, entities: Dict) -> Dict:
        """Enhance with LLM with error handling."""
        if not self.openai_client:
            logger.warning("OpenAI client not available, using fallback")
            return self._create_fallback_structure(text, entities)
        
        try:
            return self._enhance_with_llm(text, entities)
        except Exception as e:
            logger.error(f"LLM enhancement failed: {e}")
            self.warnings.append(f"LLM enhancement failed: {str(e)}")
            return self._create_fallback_structure(text, entities)
    
    def _enhance_with_llm(self, text: str, entities: Dict) -> Dict:
        """Use OpenAI to extract structured information."""
        prompt = f"""
Extract structured information from this resume and return as JSON.

Resume Text (first 6000 chars):
{text[:6000]}

Return a JSON object with:
- personal_info (name, email, phone, location, linkedin, github)
- summary (2-3 sentences)
- experience (array with: title, company, location, start_date, end_date, description)
- education (array with: degree, institution, location, graduation_date)
- skills (array of strings)
- certifications (array with: name, issuer, date)
- languages (array of strings)
- total_years_experience (number)

Return ONLY valid JSON, no other text.
"""
        
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4.1-mini",  # Use Manus-supported model
                messages=[
                    {"role": "system", "content": "You are a resume parser that outputs only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            logger.info("LLM enhancement successful")
            return result
            
        except Exception as e:
            raise LLMEnhancementError(f"OpenAI API call failed: {e}")
    
    def _create_fallback_structure(self, text: str, entities: Dict) -> Dict:
        """Create basic structure when LLM is not available."""
        logger.info("Using fallback structure")
        
        # Extract name (first person entity or first line)
        name = None
        if entities.get("persons"):
            name = entities["persons"][0]
        else:
            lines = [l.strip() for l in text.split('\n') if l.strip()]
            if lines:
                name = lines[0]
        
        return {
            "personal_info": {
                "name": name,
                "email": entities["emails"][0] if entities.get("emails") else None,
                "phone": entities["phones"][0] if entities.get("phones") else None,
                "location": entities["locations"][0] if entities.get("locations") else None,
                "linkedin": next((url for url in entities.get("urls", []) if "linkedin" in url.lower()), None),
                "github": next((url for url in entities.get("urls", []) if "github" in url.lower()), None)
            },
            "summary": None,
            "experience": [],
            "education": [],
            "skills": self._extract_skills_from_text(text),
            "certifications": [],
            "languages": [],
            "total_years_experience": None
        }
    
    def _extract_skills_from_text(self, text: str) -> List[str]:
        """Extract skills using keyword matching."""
        skill_keywords = [
            'java', 'python', 'javascript', 'react', 'angular', 'vue', 'node.js',
            'spring', 'hibernate', 'sql', 'mysql', 'postgresql', 'mongodb', 'redis',
            'aws', 'azure', 'docker', 'kubernetes', 'html', 'css', 'typescript',
            'c++', 'c#', '.net', 'django', 'flask', 'fastapi', 'git', 'jira',
            'jenkins', 'maven', 'gradle', 'rest', 'api', 'microservices'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            if skill in text_lower:
                found_skills.append(skill.title())
        
        return sorted(list(set(found_skills)))
    
    def _post_process_and_validate(self, data: Dict, text: str) -> Dict:
        """Post-process and validate extracted data."""
        # Validate email
        if data.get('personal_info', {}).get('email'):
            try:
                validate_email(data['personal_info']['email'])
            except EmailNotValidError:
                logger.warning(f"Invalid email: {data['personal_info']['email']}")
                data['personal_info']['email'] = None
        
        # Ensure required fields exist
        if 'personal_info' not in data:
            data['personal_info'] = {}
        if 'skills' not in data:
            data['skills'] = []
        if 'experience' not in data:
            data['experience'] = []
        if 'education' not in data:
            data['education'] = []
        
        # Rank and classify skills to get top 5 technology and domain skills
        if data.get('skills') and len(data['skills']) > 0:
            try:
                skill_ranking = rank_and_classify_skills(
                    resume_text=text,
                    all_skills=data['skills'],
                    experience_data=data.get('experience', []),
                    top_n=5
                )
                data['top_technology_skills'] = skill_ranking.get('top_technology_skills', [])
                data['top_domain_skills'] = skill_ranking.get('top_domain_skills', [])
                data['technology_skills_with_scores'] = skill_ranking.get('technology_skills_with_scores', [])
                data['domain_skills_with_scores'] = skill_ranking.get('domain_skills_with_scores', [])
                data['skill_ranking_metadata'] = skill_ranking.get('ranking_metadata', {})
                logger.info(f"Skill ranking complete: {len(data['top_technology_skills'])} tech, {len(data['top_domain_skills'])} domain")
            except Exception as e:
                logger.error(f"Skill ranking failed: {str(e)}")
                data['top_technology_skills'] = data['skills'][:5]
                data['top_domain_skills'] = []
        else:
            data['top_technology_skills'] = []
            data['top_domain_skills'] = []
        
        # Add success flag
        data['success'] = True
        data['errors'] = self.errors
        data['warnings'] = self.warnings
        
        return data
    
    def _create_metadata(self, file_path: str, is_scanned: bool) -> Dict:
        """Create metadata about the parsing process."""
        return {
            'file_path': file_path,
            'file_name': Path(file_path).name,
            'file_size_bytes': os.path.getsize(file_path),
            'file_extension': Path(file_path).suffix,
            'is_scanned': is_scanned,
            'parsed_at': datetime.now().isoformat(),
            'parsing_time_seconds': time.time() - self.start_time if self.start_time else 0,
            'parser_version': '2.0.0',
            'llm_used': self.openai_client is not None,
            'ner_used': self.nlp is not None
        }


# Convenience function
def parse_resume(file_path: str, openai_api_key: Optional[str] = None) -> Dict:
    """
    Parse a resume file and return structured data.
    
    Args:
        file_path: Path to the resume file
        openai_api_key: Optional OpenAI API key
        
    Returns:
        Dictionary with structured resume data
    """
    parser = EnhancedResumeParser(openai_api_key=openai_api_key)
    return parser.parse(file_path)

